"""Render del bundle de ReportService a un documento Word (.docx).

Consume el mismo `bundle` dict que antes sólo se serializaba a JSON
(`ReportService._build_bundle`): estudio, `analysis_results`, `barem_results`
(`StudyResultsOverview` con la lista de `ConstructBaremResult` — las tablas
de baremación por dimensión, ya con supresión de privacidad aplicada). No
recalcula nada: sólo formatea lo que el bundle ya trae.

Las gráficas se generan con matplotlib en backend "Agg" (sin display) y se
insertan como PNG embebido — no hay renderizado interactivo en el reporte
descargado, sólo imágenes estáticas dentro del Word.
"""

from __future__ import annotations

import io
from datetime import datetime

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402  (después de matplotlib.use)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_MUTED_RGB = RGBColor(0x6B, 0x72, 0x80)

_CLASSIFICATION_LABELS = {
    "RIESGO_ALTO": "Riesgo alto",
    "FACTOR_PROTECTOR": "Factor protector",
    "RIESGO_MEDIO": "Riesgo medio",
    "REVISION": "Requiere revisión",
}


def render_report_docx(bundle: dict) -> bytes:
    document = Document()

    study = bundle.get("study") or {}
    cover = bundle.get("cover") or {}
    traceability = bundle.get("traceability") or {}
    min_publishable_n = (traceability.get("privacy") or {}).get("min_publishable_n")
    version_kind = cover.get("version_kind")
    selected = set(bundle.get("sections") or [])

    def include(*keys: str) -> bool:
        return not selected or bool(selected.intersection(keys))

    if include("portada"):
        _add_cover(document, cover, study)
    if include("resumen_ejecutivo"):
        _add_executive_summary(document, bundle.get("executive_summary"))
    if include("trazabilidad", "ficha_tecnica"):
        _add_methodological_status(
            document, bundle.get("methodological_status"), bundle.get("methodological_label")
        )
    if include("ficha_tecnica"):
        _add_ficha_tecnica(
            document,
            study,
            bundle.get("data_quality"),
            bundle.get("methodological_label"),
            traceability.get("barem"),
            min_publishable_n,
        )
    if include("calidad_datos"):
        _add_data_quality_section(document, bundle.get("data_quality"))

    if include("variables_descriptivas"):
        _add_sociolaboral_section(document, bundle.get("sociolaboral_profile") or [])

    # Fase 4: D1-D6 y S1-S20 en bloques completamente separados — nunca una
    # tabla mixta. S1-S20 solo se renderiza para versión MEDIUM (ni título,
    # ni tabla, ni empty state en SHORT).
    if include("dimensiones", "resultados_globales"):
        document.add_heading("Resultados por dimensión (D1-D6)", level=1)
        _add_trichotomy_section(document, bundle.get("dimension_results") or [], min_publishable_n)
    if version_kind == "MEDIUM" and include("subdimensiones"):
        document.add_heading("Resultados por subdimensión (S1-S20)", level=1)
        _add_trichotomy_section(
            document, bundle.get("subdimension_results") or [], min_publishable_n, compact=True
        )

    if include("dimensiones", "resultados_globales", "subdimensiones"):
        _add_interpretation_section(document, bundle.get("interpretation") or [])
        _add_priority_ranking_section(document, bundle.get("priority_ranking") or [])

    if include("hallazgos_premium") and bundle.get("analysis_results"):
        _add_analysis_section(document, bundle["analysis_results"])
    if include("hallazgos_premium"):
        _add_premium_section(document, bundle.get("premium_analytics") or {})
    if include("plan_accion"):
        _add_action_plan_section(document, bundle.get("action_plans") or [])
    if include("conclusiones", "resumen_ejecutivo"):
        _add_conclusions_section(document, bundle.get("conclusions") or [])
    if include("anexos", "trazabilidad"):
        _add_traceability_appendix(document, traceability)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_cover(document: Document, cover: dict, study: dict) -> None:
    """Fase 1: portada/control documental. El texto "oficial"/"referencia"
    viene siempre de `censopas_methodology.py` (`cover["methodological_label"]`),
    nunca de una inferencia local."""
    title = document.add_heading(study.get("name") or "Reporte", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label = cover.get("methodological_label") or {}
    mode_label = (
        f"{label.get('label', 'Baremo de referencia')} — "
        + ("Equivalente al resultado oficial CENSOPAS-COPSOQ" if label.get("official_equivalence")
           else "No equivalente al resultado oficial CENSOPAS-COPSOQ")
    )
    subtitle = document.add_paragraph(mode_label)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.size = Pt(13)
        subtitle.runs[0].font.color.rgb = _MUTED_RGB

    date_label = _format_datetime(cover.get("generated_at"))
    if date_label:
        date_paragraph = document.add_paragraph(f"Generado el {date_label}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if date_paragraph.runs:
            date_paragraph.runs[0].font.size = Pt(10)
            date_paragraph.runs[0].font.color.rgb = _MUTED_RGB

    document.add_paragraph()
    rows = [
        ("Instrumento", cover.get("instrument_label") or "—"),
        ("Versión", cover.get("version_kind") or "UNKNOWN"),
        ("Organización", cover.get("organization_name") or "—"),
        ("Código del estudio", cover.get("study_code") or "—"),
        (
            "Período",
            " – ".join(
                filter(None, [_format_datetime(cover.get("period_start")), _format_datetime(cover.get("period_end"))])
            )
            or "—",
        ),
        ("Evaluador/responsable", cover.get("evaluator_label") or "—"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label_text, value in rows:
        cells = table.add_row().cells
        cells[0].text = label_text
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True

    confidentiality = document.add_paragraph(cover.get("confidentiality_notice") or "")
    if confidentiality.runs:
        confidentiality.runs[0].font.size = Pt(9)
        confidentiality.runs[0].font.color.rgb = _MUTED_RGB

    document.add_page_break()


def _add_methodological_status(
    document: Document, readiness: dict | None, methodological_label: dict | None
) -> None:
    if not readiness and not methodological_label:
        return
    document.add_heading("Estado metodológico", level=1)
    readiness = readiness or {}
    label = methodological_label or {}
    scoring_label = "Listo" if readiness.get("ready_for_scoring") else "Bloqueado"
    equivalence_label = (
        "Equivalente al resultado oficial CENSOPAS-COPSOQ"
        if label.get("official_equivalence")
        else "No equivalente al resultado oficial CENSOPAS-COPSOQ"
    )
    document.add_paragraph(
        f"Versión: {readiness.get('version_kind', 'UNKNOWN')}. "
        f"Scoring: {scoring_label}. "
        f"Baremo: {label.get('label', 'Baremo de referencia')}. {equivalence_label}."
    )
    if readiness.get("errors"):
        document.add_paragraph("Bloqueos: " + ", ".join(readiness["errors"]))
    if readiness.get("warnings"):
        document.add_paragraph("Advertencias: " + ", ".join(readiness["warnings"]))
    document.add_page_break()


def _add_ficha_tecnica(
    document: Document,
    study: dict,
    data_quality: dict | None,
    methodological_label: dict | None,
    barem: dict | None,
    min_publishable_n: int | None,
) -> None:
    document.add_heading("Ficha técnica", level=1)
    data_quality = data_quality or {}
    label = methodological_label or {}
    rows = [
        ("Instrumento", study.get("instrument_name") or "CENSOPAS-COPSOQ"),
        ("Versión del instrumento", study.get("instrument_version_code") or "—"),
        ("Estudio", study.get("name") or "—"),
        ("Tipo de estudio", study.get("study_type") or "—"),
        ("Estado", study.get("status") or "—"),
        ("Población convocada", "No disponible en el sistema (sin registro de convocatoria)"),
        ("Respondieron (sesiones iniciadas)", str(data_quality.get("started_count", "—"))),
        ("Válidos", str(data_quality.get("valid_count", "—"))),
        ("Excluidos", str(data_quality.get("excluded_count", "—"))),
        ("Tasa válida", _format_percentage(data_quality.get("completion_rate"))),
        (
            "Período",
            " – ".join(
                filter(
                    None,
                    [
                        _format_datetime(study.get("period_start")),
                        _format_datetime(study.get("period_end")),
                    ],
                )
            )
            or "—",
        ),
        ("Baremo aplicado", (barem or {}).get("name") or "—"),
        (
            "Equivalencia oficial",
            "Habilitada" if label.get("official_equivalence") else "No habilitada",
        ),
        ("N mínimo publicable", str(min_publishable_n) if min_publishable_n is not None else "—"),
    ]

    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label_text, value in rows:
        cells = table.add_row().cells
        cells[0].text = label_text
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
    document.add_paragraph()


def _add_executive_summary(document: Document, summary: dict | None) -> None:
    """Guía y modelos de reporte §31: primer contenido sustantivo del informe,
    antes de cualquier tabla — participación, niveles principales, prioridades."""
    if not summary:
        return
    document.add_heading("Resumen ejecutivo", level=1)
    document.add_paragraph(summary.get("headline") or "—")

    rows = [
        ("Respuestas válidas", str(summary["n_valid"]) if summary.get("n_valid") is not None else "—"),
        ("Tasa de finalización", _format_percentage(summary.get("completion_rate"))),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True

    priority_dimensions = summary.get("priority_dimensions") or []
    if priority_dimensions:
        document.add_paragraph("Dimensiones prioritarias:")
        for dimension in priority_dimensions:
            label = dimension.get("construct_name") or dimension.get("construct_code") or "—"
            classification = _CLASSIFICATION_LABELS.get(
                dimension.get("classification"), dimension.get("classification") or "—"
            )
            document.add_paragraph(
                f"{label} — {_format_pct(dimension.get('unfavorable_pct'))} "
                f"desfavorable ({classification})",
                style="List Bullet",
            )
    document.add_paragraph()


def _add_data_quality_section(document: Document, data_quality: dict | None) -> None:
    """Guía y modelos de reporte §4: calidad de captura — no es el resultado
    CENSOPAS, es participación (convocados/recibidos/válidos/tasas)."""
    if not data_quality:
        return
    document.add_heading("Calidad de datos", level=1)
    rows = [
        ("Sesiones iniciadas", str(data_quality.get("started_count", "—"))),
        ("Completadas", str(data_quality.get("completed_count", "—"))),
        ("Válidas", str(data_quality.get("valid_count", "—"))),
        ("Abandonadas", str(data_quality.get("abandoned_count", "—"))),
        ("En revisión", str(data_quality.get("review_count", "—"))),
        ("Excluidas", str(data_quality.get("excluded_count", "—"))),
        ("Tasa de finalización", _format_percentage(data_quality.get("completion_rate"))),
        ("Duración promedio", _format_duration(data_quality.get("avg_duration_seconds"))),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
    document.add_paragraph(
        "Los porcentajes se calculan sobre sesiones iniciadas; no incluyen identidad "
        "de los respondientes.",
    )
    document.add_paragraph()


_TRICHOTOMY_COLORS = {
    "favorable_pct": "2E7D32",
    "intermediate_pct": "F5B21A",
    "unfavorable_pct": "C0392B",
}
_TRICHOTOMY_LABELS = {
    "favorable_pct": "Favorable",
    "intermediate_pct": "Intermedio",
    "unfavorable_pct": "Desfavorable",
}


def _add_sociolaboral_section(document: Document, profile: list[dict]) -> None:
    """Fase 3: 11 variables C-001..C-011 (EXOGENOUS, nunca puntuables).
    Supresión n<mínimo publicable ya aplicada por `ReportService` — esta
    función solo formatea, nunca decide qué ocultar."""
    document.add_heading("Perfil sociolaboral", level=1)
    if not profile:
        document.add_paragraph(
            "Este estudio no tiene variables sociolaborales descriptivas provisionadas."
        )
        return
    document.add_paragraph(
        "Variables descriptivas (sexo, edad, instrucción, puesto, área, contrato, "
        "antigüedad, horario, horas, tipo de sueldo, rango remunerativo). No forman "
        "parte del puntaje psicosocial."
    )
    headers = ["Variable", "Categoría", "n", "%"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.bold = True
    for question in profile:
        for category in question.get("categories") or []:
            cells = table.add_row().cells
            cells[0].text = question.get("label") or question.get("code") or "—"
            cells[1].text = category.get("label") or "—"
            if category.get("suppressed"):
                cells[2].text = "Oculto"
                cells[3].text = "Oculto"
            else:
                cells[2].text = str(category.get("n", "—"))
                cells[3].text = _format_pct(category.get("percentage"))
    document.add_paragraph()


def _add_trichotomy_section(
    document: Document, results: list[dict], min_publishable_n: int | None, *, compact: bool = False
) -> None:
    """Fase 5: tabla Favorable/Intermedio/Desfavorable + n válido +
    clasificación colectiva, y barra horizontal apilada 100% — nunca pie ni
    radar. Misma fuente (`censopas_results`, trichotomy) para D1-D6 y
    S1-S20 — solo cambia qué lista se le pasa."""
    if not results:
        document.add_paragraph("Este estudio todavía no tiene resultados calculados para este bloque.")
        return

    headers = ["Constructo", "Favorable n (%)", "Intermedio n (%)", "Desfavorable n (%)", "n válido", "Clasificación colectiva"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.bold = True

    for result in results:
        cells = table.add_row().cells
        code = result.get("construct_code")
        name = result.get("construct_name") or "—"
        cells[0].text = f"{code} · {name}" if code else name

        if result.get("suppressed"):
            hidden = f"Oculto (n<{min_publishable_n})" if min_publishable_n else "Oculto"
            for i in range(1, len(headers)):
                cells[i].text = hidden
            continue

        for i, key in enumerate(("favorable_pct", "intermediate_pct", "unfavorable_pct")):
            n_key = key.replace("_pct", "_n")
            pct = result.get(key)
            n = result.get(n_key)
            cells[1 + i].text = f"{n if n is not None else '—'} ({_format_pct(pct)})"
            if pct is not None:
                _shade_cell(cells[1 + i], _TRICHOTOMY_COLORS[key])
        cells[4].text = str(result.get("n_valid", "—"))
        cells[5].text = _CLASSIFICATION_LABELS.get(
            result.get("collective_classification"), result.get("collective_classification") or "—"
        )

    document.add_paragraph()
    if not compact:
        _insert_chart(document, _build_trichotomy_chart(results))


def _add_interpretation_section(document: Document, entries: list[dict]) -> None:
    """Fase 6: hallazgo / clasificación / prioridad / hipótesis a contrastar
    / orientación preventiva / limitación, siempre separados. Nunca afirma
    causalidad ni diagnostica personas."""
    document.add_heading("Interpretación de resultados", level=1)
    if not entries:
        document.add_paragraph("Sin resultados publicables todavía para interpretar.")
        return
    for entry in entries:
        document.add_heading(entry.get("construct_name") or entry.get("construct_code") or "—", level=2)
        document.add_paragraph(f"Hallazgo: {entry.get('finding') or '—'}", style="List Bullet")
        if entry.get("classification"):
            document.add_paragraph(
                f"Clasificación: {_CLASSIFICATION_LABELS.get(entry['classification'], entry['classification'])}",
                style="List Bullet",
            )
        if entry.get("priority_rank"):
            document.add_paragraph(f"Prioridad: #{entry['priority_rank']}", style="List Bullet")
        hypotheses = entry.get("origin_hypothesis") or []
        document.add_paragraph(
            "Hipótesis de origen a contrastar: " + ("; ".join(hypotheses) if hypotheses else "sin registrar todavía."),
            style="List Bullet",
        )
        orientations = entry.get("preventive_orientation") or []
        document.add_paragraph(
            "Orientación preventiva: " + ("; ".join(orientations) if orientations else "—"),
            style="List Bullet",
        )
        document.add_paragraph(f"Limitación: {entry.get('limitation') or '—'}", style="List Bullet")
    document.add_paragraph()


def _add_priority_ranking_section(document: Document, ranking: list[dict]) -> None:
    """Fase 7: orden de priorización preventiva — nunca una clasificación
    oficial adicional."""
    document.add_heading("Priorización preventiva", level=1)
    document.add_paragraph(
        "El siguiente orden es una ayuda de priorización preventiva calculada por el "
        "backend (% desfavorable); no es una clasificación oficial adicional de "
        "CENSOPAS-COPSOQ ni reemplaza el juicio técnico del Comité/Grupo de Trabajo."
    )
    if not ranking:
        document.add_paragraph("Sin dimensiones publicables para priorizar todavía.")
        return
    headers = ["Prioridad", "Constructo", "% Desfavorable", "n válido", "Clasificación colectiva"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in ranking:
        cells = table.add_row().cells
        cells[0].text = str(row.get("priority_rank") or "—")
        cells[1].text = f"{row.get('construct_code') or ''} · {row.get('construct_name') or '—'}"
        cells[2].text = _format_pct(row.get("unfavorable_pct"))
        cells[3].text = str(row.get("n_valid") or "—")
        cells[4].text = _CLASSIFICATION_LABELS.get(
            row.get("collective_classification"), row.get("collective_classification") or "—"
        )
    document.add_paragraph()


def _add_conclusions_section(document: Document, conclusions: list[str]) -> None:
    """Fase 9: solo texto ya derivado en `ReportService._build_conclusions` —
    esta función únicamente formatea la lista."""
    document.add_heading("Conclusiones", level=1)
    if not conclusions:
        document.add_paragraph("Sin datos suficientes para conclusiones todavía.")
        return
    for line in conclusions:
        document.add_paragraph(line, style="List Bullet")
    document.add_paragraph()


def _add_analysis_section(document: Document, analysis_results: list[dict]) -> None:
    document.add_heading("Resultados de análisis estadístico", level=1)
    headers = ["Código", "Tipo", "N válido", "Valor", "Estadístico", "p", "Tamaño de efecto"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True

    # Cap defensivo: un reporte con cientos de result_type sueltos no debería
    # convertirse en un Word de mil filas — la sección de baremación es la
    # tabla protagonista, ésta es un complemento.
    for result in analysis_results[:50]:
        cells = table.add_row().cells
        cells[0].text = result.get("result_code") or "—"
        cells[1].text = result.get("result_type") or "—"
        cells[2].text = str(result["n_valid"]) if result.get("n_valid") is not None else "—"
        cells[3].text = _format_number(result.get("numeric_value"))
        cells[4].text = _format_number(result.get("statistic_value"))
        cells[5].text = _format_number(result.get("p_value"), decimals=4)
        cells[6].text = _format_number(result.get("effect_size"))


def _add_action_plan_section(document: Document, plans: list[dict]) -> None:
    document.add_heading("Plan de acción y seguimiento", level=1)
    if not plans:
        document.add_paragraph("No hay planes de acción registrados para este estudio.")
        return
    for plan in plans:
        document.add_heading(plan.get("name") or "Plan de acción", level=2)
        approval = "Aprobado" if plan.get("approved") else "Pendiente de aprobación"
        document.add_paragraph(f"Estado: {plan.get('status') or '—'}. {approval}.")
        headers = [
            "Prioridad",
            "Constructo",
            "Hallazgo",
            "Hipótesis de origen",
            "Medida",
            "Responsable",
            "Fecha",
            "Estado",
        ]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in plan.get("items") or []:
            cells = table.add_row().cells
            cells[0].text = str(item.get("priority") or "—")
            construct = item.get("construct_code") or "—"
            if item.get("construct_name"):
                construct += f" · {item['construct_name']}"
            cells[1].text = construct
            cells[2].text = item.get("finding") or "—"
            cells[3].text = item.get("origin_hypothesis") or "—"
            cells[4].text = item.get("action_description") or "—"
            cells[5].text = item.get("responsible_label") or "—"
            cells[6].text = item.get("due_date") or "—"
            cells[7].text = item.get("effective_status") or item.get("status") or "—"
            for kpi in item.get("kpis") or []:
                measured = kpi.get("current_value")
                if measured is None:
                    latest = kpi.get("latest_measurement") or {}
                    measured = latest.get("text_value") or "sin medición"
                document.add_paragraph(
                    f"KPI {kpi.get('code') or '—'}: {kpi.get('name') or '—'} · "
                    f"línea base {_format_number(kpi.get('baseline_value'))} · "
                    f"actual {measured} · "
                    f"meta {_format_number(kpi.get('target_value'))}."
                )
        document.add_paragraph()


def _add_premium_section(document: Document, premium: dict) -> None:
    document.add_heading("Analítica premium", level=1)
    if premium.get("status") != "AVAILABLE":
        document.add_paragraph("No hay resultados premium publicables para este estudio.")
        return
    document.add_paragraph(
        f"{premium.get('result_count', 0)} resultados publicables; "
        f"{premium.get('significant_result_count', 0)} con p ajustado o p menor a 0.05. "
        f"Métodos: {', '.join(premium.get('methods') or []) or '—'}."
    )
    headers = ["Código", "Método", "N", "Estadístico", "p ajustado", "Efecto", "IC"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for result in (premium.get("results") or [])[:100]:
        cells = table.add_row().cells
        cells[0].text = result.get("result_code") or "—"
        cells[1].text = result.get("result_type") or "—"
        cells[2].text = str(result.get("n_valid") if result.get("n_valid") is not None else "—")
        cells[3].text = _format_number(result.get("statistic_value"))
        adjusted = result.get("adjusted_p_value")
        cells[4].text = _format_number(adjusted if adjusted is not None else result.get("p_value"), 4)
        cells[5].text = " · ".join(filter(None, [_format_number(result.get("effect_size")), result.get("effect_label")]))
        cells[6].text = f"[{_format_number(result.get('ci_lower'))}, {_format_number(result.get('ci_upper'))}]"
    document.add_paragraph("Limitaciones: " + " ".join(premium.get("limitations") or []))


def _add_traceability_appendix(document: Document, traceability: dict) -> None:
    document.add_heading("Anexo de trazabilidad", level=1)
    if not traceability:
        document.add_paragraph("No hay metadatos de trazabilidad disponibles.")
        return
    rows = [
        ("Versión del instrumento", traceability.get("version_kind") or "—"),
        ("Hash del manifiesto", traceability.get("manifest_hash") or "—"),
        ("Linaje", " → ".join(traceability.get("lineage") or [])),
        ("N mínimo publicable", str((traceability.get("privacy") or {}).get("min_publishable_n", "—"))),
        ("Registros individuales incluidos", "No"),
    ]
    barem = traceability.get("barem") or {}
    if barem:
        rows.extend([
            ("Baremo", f"{barem.get('name') or '—'} · {barem.get('version') or '—'}"),
            ("Hash del baremo", barem.get("content_hash") or "—"),
            ("Fuente del baremo", barem.get("source_reference") or "—"),
        ])
    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    runs = traceability.get("analysis_runs") or []
    if runs:
        document.add_heading("Ejecuciones analíticas", level=2)
        run_table = document.add_table(rows=1, cols=5)
        run_table.style = "Light Grid Accent 1"
        for index, header in enumerate(["Tipo", "Estado", "Motor", "Algoritmo", "Hash de entrada"]):
            run_table.rows[0].cells[index].text = header
        for run in runs:
            cells = run_table.add_row().cells
            cells[0].text = run.get("analysis_type") or "—"
            cells[1].text = run.get("status") or "—"
            cells[2].text = " ".join(filter(None, [run.get("engine"), run.get("engine_version")])) or "—"
            cells[3].text = run.get("algorithm_version") or "—"
            cells[4].text = run.get("input_hash") or "—"


def _build_trichotomy_chart(results: list[dict]) -> io.BytesIO | None:
    """Fase 5: barra horizontal apilada 100% Favorable/Intermedio/Desfavorable
    — nunca pie ni radar. Una barra por constructo, en el orden recibido
    (ya viene ordenado por `priority_rank`)."""
    plot_results = [r for r in results if not r.get("suppressed") and r.get("unfavorable_pct") is not None]
    if not plot_results:
        return None

    labels = [r.get("construct_code") or r.get("construct_name") or "" for r in plot_results]
    fig_height = max(2.2, 0.5 * len(plot_results) + 1.3)
    fig, ax = plt.subplots(figsize=(6.3, fig_height), dpi=150)

    y_pos = list(range(len(plot_results)))
    left = [0.0] * len(plot_results)
    for key in ("favorable_pct", "intermediate_pct", "unfavorable_pct"):
        widths = [r.get(key) or 0 for r in plot_results]
        ax.barh(
            y_pos, widths, left=left, color=f"#{_TRICHOTOMY_COLORS[key]}",
            label=_TRICHOTOMY_LABELS[key], edgecolor="white", height=0.62,
        )
        left = [total + width for total, width in zip(left, widths)]

    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("% de encuestados", fontsize=9)
    ax.set_title("Distribución favorable / intermedio / desfavorable", fontsize=11, fontweight="bold")
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.15), ncol=3, fontsize=8, frameon=False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _fig_to_png(fig) -> io.BytesIO:
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _insert_chart(document: Document, image: io.BytesIO | None) -> None:
    if image is None:
        return
    document.add_picture(image, width=Cm(15))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _format_number(value: float | None, decimals: int = 2) -> str:
    if value is None:
        return "—"
    return f"{value:.{decimals}f}"


def _format_pct(value: float | None, decimals: int = 1) -> str:
    """Formatea un valor YA en escala 0-100 (ej. `unfavorable_pct`)."""
    if value is None:
        return "—"
    return f"{value:.{decimals}f}%"


def _format_percentage(ratio: float | None, decimals: int = 1) -> str:
    """Formatea una razón 0-1 (ej. `completion_rate`) como porcentaje."""
    if ratio is None:
        return "—"
    return f"{ratio * 100:.{decimals}f}%"


def _format_duration(seconds: float | None) -> str:
    if seconds is None:
        return "—"
    minutes, remainder = divmod(int(seconds), 60)
    return f"{minutes} min {remainder} s" if minutes else f"{remainder} s"


def _format_datetime(value: str | None) -> str | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value).strftime("%d/%m/%Y %H:%M")
    except ValueError:
        return value
