"""Widgets compuestos: header/footer editorial, badges, títulos de sección,
callouts y la grilla de firmas. El renderer de secciones sólo llama a estas
funciones — nunca manipula OOXML directamente."""

from __future__ import annotations

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from .borders import add_horizontal_rule, remove_table_borders, set_cell_border, set_cell_margins
from .cells import center_table, set_cell_vertical_center, set_column_widths_mm, set_table_fixed_layout, shade_cell
from .paginate import keep_table_row_together
from .theme import COLMENA, CONTENT_WIDTH_MM, FONT, PAGE
from .typography import add_h1

_HAIRLINE = {"sz": 3, "color": COLMENA["line"], "val": "single"}


def _run(paragraph, text: str, *, size: float, bold: bool = False, color: str = COLMENA["text"]):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_separate, fld_end):
        run._r.append(element)
    run.font.size = Pt(FONT["caption"])
    run.font.color.rgb = RGBColor.from_string(COLMENA["muted"])


def add_header_band(document) -> None:
    """Franja navy con el nombre de marca y una línea dorada parcial debajo
    — nunca debe parecer una tabla.

    El bleed literal a los bordes físicos de página (`tblInd` negativo) no
    se sostiene de forma fiable en el renderer real del pipeline (LibreOffice
    headless vía `report_pdf_convert.docx_to_pdf`): el frame de header/footer
    de LO está acotado al área entre márgenes y una tabla más ancha que ese
    frame se recorta/renderiza mal. La franja ocupa entonces todo el ancho de
    contenido (margen a margen) — mismo peso visual, sin depender de un
    comportamiento de bleed que el propio pipeline no reproduce."""
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for paragraph in header.paragraphs:
        paragraph.text = ""

    band = header.add_table(rows=1, cols=1, width=Mm(CONTENT_WIDTH_MM))
    set_table_fixed_layout(band, CONTENT_WIDTH_MM)
    set_column_widths_mm(band, [CONTENT_WIDTH_MM])
    band.rows[0].height = Mm(10)
    cell = band.rows[0].cells[0]
    shade_cell(cell, COLMENA["navy"])
    set_cell_margins(cell, top=90, bottom=90, start=int(Mm(6).twips), end=100)
    set_cell_vertical_center(cell)
    remove_table_borders(band)
    paragraph = cell.paragraphs[0]
    _run(paragraph, "COLMENA | INTELIGENCIA PSICOSOCIAL", size=8.5, bold=True, color=COLMENA["white"])

    # Word/LibreOffice tratan dos `w:tbl` consecutivas sin un párrafo entre
    # medio como una única tabla fusionada (la grilla de la segunda termina
    # aplicándose a la primera) — este párrafo separador de 0pt es
    # obligatorio, no cosmético.
    spacer = header.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    if spacer.runs:
        spacer.runs[0].font.size = Pt(1)

    rule = header.add_table(rows=1, cols=1, width=Mm(40))
    set_table_fixed_layout(rule, 40)
    set_column_widths_mm(rule, [40])
    rule.rows[0].height = Mm(1.3)
    rule_cell = rule.rows[0].cells[0]
    shade_cell(rule_cell, COLMENA["gold"])
    remove_table_borders(rule)


def add_footer_band(document) -> None:
    """Regla fina + identificación a la izquierda + paginación a la derecha
    — dos columnas 85/15, nunca un párrafo centrado."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for paragraph in footer.paragraphs:
        paragraph.text = ""

    content_width_mm = PAGE["width"] / Mm(1) - PAGE["margin_left"] / Mm(1) - PAGE["margin_right"] / Mm(1)
    left_width = content_width_mm * 0.85
    right_width = content_width_mm * 0.15

    table = footer.add_table(rows=1, cols=2, width=Mm(content_width_mm))
    set_table_fixed_layout(table, content_width_mm)
    remove_table_borders(table)
    set_cell_border(table.rows[0].cells[0], top=_HAIRLINE)
    set_cell_border(table.rows[0].cells[1], top=_HAIRLINE)

    left_cell, right_cell = table.rows[0].cells
    set_cell_margins(left_cell, top=60, bottom=0, start=0, end=60)
    set_cell_margins(right_cell, top=60, bottom=0, start=60, end=0)
    set_cell_vertical_center(left_cell)
    set_cell_vertical_center(right_cell)

    left_paragraph = left_cell.paragraphs[0]
    _run(left_paragraph, "CENSOPAS-COPSOQ · expediente agregado · privacidad protegida", size=FONT["caption"], color=COLMENA["muted"])

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(right_paragraph, "Página ", size=FONT["caption"], color=COLMENA["muted"])
    add_page_field(right_paragraph, "PAGE")

    set_column_widths_mm(table, [left_width, right_width])


def add_centered_badge(document, text: str, *, width_mm: float = 70, height_mm: float = 9, background: str = COLMENA["gold"], color: str = COLMENA["white"], size: float = 8.5) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_fixed_layout(table, width_mm)
    center_table(table)
    table.rows[0].height = Mm(height_mm)
    cell = table.rows[0].cells[0]
    shade_cell(cell, background)
    set_cell_vertical_center(cell)
    remove_table_borders(table)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(paragraph, text, size=size, bold=True, color=color)
    set_column_widths_mm(table, [width_mm])


def add_section_header(document, number: str, title: str) -> None:
    """Badge turquesa centrado + título 20-22pt + regla dorada. Único
    patrón de encabezado de sección en todo el documento."""
    document.add_paragraph().paragraph_format.space_after = Pt(6)
    add_centered_badge(document, number, width_mm=16, height_mm=9, background=COLMENA["teal"], size=9.5)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    heading = add_h1(document, title)
    rule_paragraph = document.add_paragraph()
    rule_paragraph.paragraph_format.space_after = Pt(8)
    add_horizontal_rule(rule_paragraph, color=COLMENA["gold"], size=8)
    return heading


_CALLOUT_VARIANTS = {
    "neutral": {"bg": COLMENA["surface"], "border": COLMENA["line"]},
    "warning": {"bg": "FBF1DC", "border": COLMENA["gold"]},
    "danger": {"bg": "FBE6E6", "border": COLMENA["red"]},
    "success": {"bg": "E3F3EC", "border": COLMENA["green"]},
}


def add_callout(document, text: str, *, variant: str = "neutral") -> None:
    palette = _CALLOUT_VARIANTS.get(variant, _CALLOUT_VARIANTS["neutral"])
    table = document.add_table(rows=1, cols=1)
    from .theme import CONTENT_WIDTH_MM

    set_table_fixed_layout(table, CONTENT_WIDTH_MM)
    cell = table.rows[0].cells[0]
    shade_cell(cell, palette["bg"])
    set_cell_margins(cell, top=140, bottom=140, start=140, end=140)
    line_border = {"sz": 4, "color": palette["border"], "val": "single"}
    set_cell_border(cell, top=line_border, bottom=line_border, left=line_border, right=line_border)
    paragraph = cell.paragraphs[0]
    _run(paragraph, text, size=FONT["body"], color=COLMENA["text"])
    for run in paragraph.runs:
        run.font.size = Pt(FONT["body"])
    set_column_widths_mm(table, [CONTENT_WIDTH_MM])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_signature_grid(document, entries: list[dict]) -> None:
    """Disposición 2x2: línea de firma, nombre, rol, registro — nunca una
    tabla Rol/Firma/Fecha de tres columnas."""
    from .theme import CONTENT_WIDTH_MM

    padded = list(entries) + [None] * (4 - len(entries)) if len(entries) < 4 else entries[:4]
    n_rows = 2
    table = document.add_table(rows=n_rows, cols=2)
    set_table_fixed_layout(table, CONTENT_WIDTH_MM)
    box_border = {"sz": 4, "color": COLMENA["line"], "val": "single"}
    set_cell_border(table.rows[0].cells[0], top=box_border, left=box_border, right=box_border, bottom=box_border)
    set_cell_border(table.rows[0].cells[1], top=box_border, left=box_border, right=box_border, bottom=box_border)
    set_cell_border(table.rows[1].cells[0], top=box_border, left=box_border, right=box_border, bottom=box_border)
    set_cell_border(table.rows[1].cells[1], top=box_border, left=box_border, right=box_border, bottom=box_border)
    for row in table.rows:
        keep_table_row_together(row)

    for index, entry in enumerate(padded):
        row, col = divmod(index, 2)
        cell = table.rows[row].cells[col]
        set_cell_margins(cell, top=280, bottom=200, start=200, end=200)
        cell.paragraphs[0].text = ""
        line_paragraph = cell.paragraphs[0]
        line_paragraph.paragraph_format.space_before = Pt(28)
        line_paragraph.paragraph_format.space_after = Pt(2)
        add_horizontal_rule(line_paragraph, color=COLMENA["muted"], size=4)
        if entry:
            name_p = cell.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(name_p, entry.get("name") or "—", size=FONT["body"], bold=True, color=COLMENA["text"])
            for key in ("role", "registry"):
                if entry.get(key):
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(p, entry[key], size=FONT["caption"], color=COLMENA["muted"])
            caption_p = cell.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(caption_p, "Firma y fecha", size=FONT["caption"], color=COLMENA["muted"])
        else:
            for _ in range(3):
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, "—", size=FONT["caption"], color=COLMENA["muted"])

    set_column_widths_mm(table, [CONTENT_WIDTH_MM / 2, CONTENT_WIDTH_MM / 2])
