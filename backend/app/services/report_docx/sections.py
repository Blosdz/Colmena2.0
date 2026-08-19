"""Contenido del reporte CENSOPAS — consume el `bundle` de
`ReportService._build_bundle` y sólo formatea: nunca recalcula nada. Cada
función de sección compone el layout llamando a `components`/`tables`/
`typography`; nunca toca color, borde o padding directamente (regla #22)."""

from __future__ import annotations

import io
from datetime import datetime

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .charts import build_trichotomy_chart
from .components import add_callout, add_centered_badge, add_section_header, add_signature_grid
from .paginate import keep_with_next, page_break_before
from .tables import add_colmena_table, add_key_value_grid, add_label_value_table
from .theme import COLMENA, CONTENT_WIDTH_MM, FONT
from .typography import add_body, add_cover_org, add_cover_subtitle, add_cover_title, add_h1, add_h2, add_h3, add_caption, add_muted

_CLASSIFICATION_LABELS = {
    "RIESGO_ALTO": "Riesgo alto",
    "FACTOR_PROTECTOR": "Factor protector",
    "RIESGO_MEDIO": "Riesgo medio",
    "REVISION": "Requiere revisión",
}
_CLASSIFICATION_STATUS = {
    "RIESGO_ALTO": "red",
    "FACTOR_PROTECTOR": "green",
    "RIESGO_MEDIO": "gold",
    "REVISION": "gold",
}
_ACTION_STATUS = {
    "Aprobado": "green",
    "Completado": "green",
    "DONE": "green",
    "COMPLETED": "green",
    "Pendiente": "gold",
    "PENDING": "gold",
    "En progreso": "gold",
    "IN_PROGRESS": "gold",
    "Vencido": "red",
    "OVERDUE": "red",
    "Bloqueado": "red",
    "BLOCKED": "red",
}


def build_document_content(document, bundle: dict) -> None:
    numbering = _section_numbering()

    study = bundle.get("study") or {}
    cover = bundle.get("cover") or {}
    traceability = bundle.get("traceability") or {}
    min_publishable_n = (traceability.get("privacy") or {}).get("min_publishable_n")
    version_kind = cover.get("version_kind")
    selected = set(bundle.get("sections") or [])

    def include(*keys: str) -> bool:
        return not selected or bool(selected.intersection(keys))

    if include("portada"):
        _add_cover(document, cover, study)
    if include("resumen_ejecutivo"):
        _add_executive_summary(document, next(numbering), bundle.get("executive_summary"))
    if include("trazabilidad", "ficha_tecnica"):
        _add_methodological_status(document, bundle.get("methodological_status"), bundle.get("methodological_label"))
    if include("ficha_tecnica"):
        _add_ficha_tecnica(
            document, next(numbering), study, bundle.get("data_quality"),
            bundle.get("methodological_label"), traceability.get("barem"), min_publishable_n,
        )
    if include("calidad_datos"):
        _add_data_quality_section(document, next(numbering), bundle.get("data_quality"))
    if include("variables_descriptivas"):
        _add_sociolaboral_section(document, next(numbering), bundle.get("sociolaboral_profile") or [])

    if include("dimensiones", "resultados_globales"):
        add_section_header(document, next(numbering), "Resultados por dimensión (D1-D6)")
        _add_trichotomy_section(document, bundle.get("dimension_results") or [], min_publishable_n)
    if version_kind == "MEDIUM" and include("subdimensiones"):
        add_section_header(document, next(numbering), "Resultados por subdimensión (S1-S20)")
        _add_trichotomy_section(document, bundle.get("subdimension_results") or [], min_publishable_n, compact=True)

    if include("dimensiones", "resultados_globales", "subdimensiones"):
        _add_interpretation_section(document, bundle.get("interpretation") or [])
        _add_priority_ranking_section(document, next(numbering), bundle.get("priority_ranking") or [])

    if include("hallazgos_premium") and (bundle.get("analysis_results") or bundle.get("premium_analytics")):
        add_section_header(document, next(numbering), "Analítica")
    if include("hallazgos_premium") and bundle.get("analysis_results"):
        _add_analysis_section(document, bundle["analysis_results"])
    if include("hallazgos_premium"):
        _add_premium_section(document, bundle.get("premium_analytics") or {})
    if include("plan_accion"):
        _add_action_plan_section(document, next(numbering), bundle.get("action_plans") or [])
    if include("conclusiones", "resumen_ejecutivo"):
        _add_conclusions_section(document, bundle.get("conclusions") or [])
    if include("anexos", "trazabilidad"):
        _add_traceability_appendix(document, next(numbering), traceability)
    if include("firmas", "anexos"):
        _add_signatures_section(document, next(numbering))


def _section_numbering():
    n = 0
    while True:
        n += 1
        yield f"{n:02d}"


# ---------------------------------------------------------------------------
# Portada
# ---------------------------------------------------------------------------

def _add_cover(document, cover: dict, study: dict) -> None:
    add_centered_badge(document, "EXPEDIENTE TÉCNICO · BORRADOR CONTROLADO", width_mm=78, height_mm=9)
    document.add_paragraph().paragraph_format.space_after = Pt(10)

    add_cover_title(document, study.get("name") or "Reporte")

    org = cover.get("organization_name")
    if org:
        add_cover_org(document, org)

    label = cover.get("methodological_label") or {}
    mode_label = (
        f"{label.get('label', 'Baremo de referencia')} — "
        + ("Equivalente al resultado oficial CENSOPAS-COPSOQ" if label.get("official_equivalence")
           else "Reporte provisional — datos sintéticos y sin equivalencia oficial")
    )
    add_cover_subtitle(document, mode_label)

    pairs = [
        ("RUC", cover.get("ruc") or study.get("ruc") or "—"),
        ("Población", str(cover.get("population") or study.get("population") or "—")),
        ("Actividad", cover.get("activity") or "—"),
        ("Versión", cover.get("version_kind") or "UNKNOWN"),
        ("Estado", study.get("status") or "—"),
        ("Privacidad", "Sin respuestas individuales · n mínimo protegido"),
    ]
    add_key_value_grid(document, pairs, widths_mm=[28, 61, 28, 61])
    document.add_paragraph().paragraph_format.space_after = Pt(14)

    add_callout(
        document,
        cover.get("confidentiality_notice")
        or "Documento para revisión del profesional responsable, psicología ocupacional, "
        "seguridad y representante de la empresa. No establece diagnósticos individuales "
        "ni inferencias causales.",
        variant="neutral",
    )

    date_label = _format_datetime(cover.get("generated_at"))
    if date_label:
        document.add_paragraph()
        add_caption(document, f"Generado: {date_label}")

    document.add_page_break()


# ---------------------------------------------------------------------------
# Estado metodológico / ficha técnica
# ---------------------------------------------------------------------------

def _add_methodological_status(document, readiness: dict | None, methodological_label: dict | None) -> None:
    if not readiness and not methodological_label:
        return
    add_h2(document, "Estado metodológico")
    readiness = readiness or {}
    label = methodological_label or {}
    scoring_label = "Listo" if readiness.get("ready_for_scoring") else "Bloqueado"
    equivalence_label = (
        "Equivalente al resultado oficial CENSOPAS-COPSOQ"
        if label.get("official_equivalence")
        else "No equivalente al resultado oficial CENSOPAS-COPSOQ"
    )
    add_body(
        document,
        f"Versión: {readiness.get('version_kind', 'UNKNOWN')}. Scoring: {scoring_label}. "
        f"Baremo: {label.get('label', 'Baremo de referencia')}. {equivalence_label}.",
    )
    if readiness.get("errors"):
        add_callout(document, "Bloqueos: " + ", ".join(readiness["errors"]), variant="danger")
    if readiness.get("warnings"):
        add_callout(document, "Advertencias: " + ", ".join(readiness["warnings"]), variant="warning")
    document.add_page_break()


def _add_ficha_tecnica(document, number, study, data_quality, methodological_label, barem, min_publishable_n) -> None:
    add_section_header(document, number, "Ficha técnica")
    data_quality = data_quality or {}
    label = methodological_label or {}
    period = " – ".join(filter(None, [_format_datetime(study.get("period_start")), _format_datetime(study.get("period_end"))])) or "—"
    pairs = [
        ("Instrumento", study.get("instrument_name") or "CENSOPAS-COPSOQ"),
        ("Versión del instrumento", study.get("instrument_version_code") or "—"),
        ("Estudio", study.get("name") or "—"),
        ("Tipo de estudio", study.get("study_type") or "—"),
        ("Estado", study.get("status") or "—"),
        ("Población convocada", "No disponible en el sistema"),
        ("Respondieron (sesiones)", str(data_quality.get("started_count", "—"))),
        ("Válidos", str(data_quality.get("valid_count", "—"))),
        ("Excluidos", str(data_quality.get("excluded_count", "—"))),
        ("Tasa válida", _format_percentage(data_quality.get("completion_rate"))),
        ("Período", period),
        ("Baremo aplicado", (barem or {}).get("name") or "—"),
        ("Equivalencia oficial", "Habilitada" if label.get("official_equivalence") else "No habilitada"),
        ("N mínimo publicable", str(min_publishable_n) if min_publishable_n is not None else "—"),
    ]
    add_key_value_grid(document, pairs, widths_mm=[38, 51, 38, 51])
    document.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_executive_summary(document, number, summary: dict | None) -> None:
    if not summary:
        return
    add_section_header(document, number, "Resumen ejecutivo")
    add_body(document, summary.get("headline") or "—")

    pairs = [
        ("Respuestas válidas", str(summary["n_valid"]) if summary.get("n_valid") is not None else "—"),
        ("Tasa de finalización", _format_percentage(summary.get("completion_rate"))),
    ]
    add_key_value_grid(document, pairs, widths_mm=[45, 44, 45, 44])

    priority_dimensions = summary.get("priority_dimensions") or []
    if priority_dimensions:
        document.add_paragraph().paragraph_format.space_after = Pt(4)
        add_h3(document, "Dimensiones prioritarias")
        rows = []
        for dimension in priority_dimensions:
            label = dimension.get("construct_name") or dimension.get("construct_code") or "—"
            classification = dimension.get("classification")
            rows.append([label, _format_pct(dimension.get("unfavorable_pct")), _CLASSIFICATION_LABELS.get(classification, classification or "—")])
        add_colmena_table(document, ["Dimensión", "% Desfavorable", "Clasificación"], rows, widths_mm=[88, 45, 45])
    document.add_paragraph()


def _add_data_quality_section(document, number, data_quality: dict | None) -> None:
    if not data_quality:
        return
    add_section_header(document, number, "Calidad de datos")
    pairs = [
        ("Sesiones iniciadas", str(data_quality.get("started_count", "—"))),
        ("Completadas", str(data_quality.get("completed_count", "—"))),
        ("Válidas", str(data_quality.get("valid_count", "—"))),
        ("Abandonadas", str(data_quality.get("abandoned_count", "—"))),
        ("En revisión", str(data_quality.get("review_count", "—"))),
        ("Excluidas", str(data_quality.get("excluded_count", "—"))),
        ("Tasa de finalización", _format_percentage(data_quality.get("completion_rate"))),
        ("Duración promedio", _format_duration(data_quality.get("avg_duration_seconds"))),
    ]
    add_key_value_grid(document, pairs, widths_mm=[38, 51, 38, 51])
    document.add_paragraph().paragraph_format.space_after = Pt(4)
    add_caption(document, "Los porcentajes se calculan sobre sesiones iniciadas; no incluyen identidad de los respondientes.")
    document.add_paragraph()


def _add_sociolaboral_section(document, number, profile: list[dict]) -> None:
    add_section_header(document, number, "Perfil sociolaboral")
    if not profile:
        add_body(document, "Este estudio no tiene variables sociolaborales descriptivas provisionadas.")
        return
    add_body(
        document,
        "Variables descriptivas (sexo, edad, instrucción, puesto, área, contrato, antigüedad, "
        "horario, horas, tipo de sueldo, rango remunerativo). No forman parte del puntaje psicosocial.",
    )
    rows = []
    for question in profile:
        for category in question.get("categories") or []:
            label = question.get("label") or question.get("code") or "—"
            cat_label = category.get("label") or "—"
            if category.get("suppressed"):
                rows.append([label, cat_label, "Oculto", "Oculto"])
            else:
                rows.append([label, cat_label, str(category.get("n", "—")), _format_pct(category.get("percentage"))])
    add_colmena_table(document, ["Variable", "Categoría", "n", "%"], rows, widths_mm=[54, 70, 27, 27], compact=True)
    document.add_paragraph()


# ---------------------------------------------------------------------------
# Resultados D1-D6 / S1-S20
# ---------------------------------------------------------------------------

def _add_trichotomy_section(document, results: list[dict], min_publishable_n, *, compact: bool = False) -> None:
    if not results:
        add_body(document, "Este estudio todavía no tiene resultados calculados para este bloque.")
        return

    headers = ["Constructo", "Favorable n (%)", "Intermedio n (%)", "Desfavorable n (%)", "n válido", "Clasificación"]
    widths = [46, 26, 26, 26, 20, 34]
    rows = []
    for result in results:
        code = result.get("construct_code")
        name = result.get("construct_name") or "—"
        label = f"{code} · {name}" if code else name
        if result.get("suppressed"):
            hidden = f"Oculto (n<{min_publishable_n})" if min_publishable_n else "Oculto"
            rows.append([label, hidden, hidden, hidden, hidden, hidden])
            continue
        cells = [label]
        for key in ("favorable_pct", "intermediate_pct", "unfavorable_pct"):
            n_key = key.replace("_pct", "_n")
            pct = result.get(key)
            n = result.get(n_key)
            cells.append(f"{n if n is not None else '—'} ({_format_pct(pct)})")
        cells.append(str(result.get("n_valid", "—")))
        classification = result.get("collective_classification")
        cells.append(_CLASSIFICATION_LABELS.get(classification, classification or "—"))
        rows.append(cells)

    table = add_colmena_table(document, headers, rows, widths_mm=widths, compact=compact)
    for row_index, result in enumerate(results):
        classification = result.get("collective_classification")
        status = _CLASSIFICATION_STATUS.get(classification)
        if status and not result.get("suppressed"):
            cell = table.rows[row_index + 1].cells[5]
            label = _CLASSIFICATION_LABELS.get(classification, classification or "—")
            _render_status(cell, label, status)

    document.add_paragraph()
    if not compact:
        _insert_chart(document, build_trichotomy_chart(results))


def _render_status(cell, label: str, status: str) -> None:
    from .cells import render_status_cell

    render_status_cell(cell, label=label, status=status)


def _add_interpretation_section(document, entries: list[dict]) -> None:
    add_h2(document, "Interpretación de resultados")
    if not entries:
        add_body(document, "Sin resultados publicables todavía para interpretar.")
        return
    for entry in entries:
        heading = add_h3(document, entry.get("construct_name") or entry.get("construct_code") or "—")
        keep_with_next(heading)
        add_body(document, f"Hallazgo: {entry.get('finding') or '—'}")
        if entry.get("classification"):
            add_body(document, f"Clasificación: {_CLASSIFICATION_LABELS.get(entry['classification'], entry['classification'])}")
        if entry.get("priority_rank"):
            add_body(document, f"Prioridad: #{entry['priority_rank']}")
        hypotheses = entry.get("origin_hypothesis") or []
        add_body(document, "Hipótesis de origen a contrastar: " + ("; ".join(hypotheses) if hypotheses else "sin registrar todavía."))
        orientations = entry.get("preventive_orientation") or []
        add_body(document, "Orientación preventiva: " + ("; ".join(orientations) if orientations else "—"))
        add_muted(document, f"Limitación: {entry.get('limitation') or '—'}")
    document.add_paragraph()


def _add_priority_ranking_section(document, number, ranking: list[dict]) -> None:
    add_section_header(document, number, "Priorización preventiva")
    add_body(
        document,
        "El siguiente orden es una ayuda de priorización preventiva calculada por el backend "
        "(% desfavorable); no es una clasificación oficial adicional de CENSOPAS-COPSOQ ni "
        "reemplaza el juicio técnico del Comité/Grupo de Trabajo.",
    )
    if not ranking:
        add_body(document, "Sin dimensiones publicables para priorizar todavía.")
        return
    rows = []
    for row in ranking:
        classification = row.get("collective_classification")
        rows.append([
            str(row.get("priority_rank") or "—"),
            f"{row.get('construct_code') or ''} · {row.get('construct_name') or '—'}",
            _format_pct(row.get("unfavorable_pct")),
            str(row.get("n_valid") or "—"),
            _CLASSIFICATION_LABELS.get(classification, classification or "—"),
        ])
    table = add_colmena_table(document, ["Prioridad", "Constructo", "% Desfavorable", "n válido", "Clasificación"], rows, widths_mm=[22, 60, 32, 24, 40])
    for row_index, row in enumerate(ranking):
        status = _CLASSIFICATION_STATUS.get(row.get("collective_classification"))
        if status:
            label = _CLASSIFICATION_LABELS.get(row.get("collective_classification"), "—")
            _render_status(table.rows[row_index + 1].cells[4], label, status)
    document.add_paragraph()


def _add_conclusions_section(document, conclusions: list[str]) -> None:
    add_h2(document, "Conclusiones")
    if not conclusions:
        add_body(document, "Sin datos suficientes para conclusiones todavía.")
        return
    for line in conclusions:
        document.add_paragraph(line, style="List Bullet")
    document.add_paragraph()


_ANALYSIS_CATEGORIES = [
    ("Confiabilidad", {"CRONBACH_ALPHA", "MCDONALD_OMEGA"}),
    ("Normalidad", {"NORMALITY"}),
    ("Comparaciones entre grupos", {"CHI_SQUARE", "MANN_WHITNEY", "KRUSKAL_WALLIS"}),
    ("Correlaciones", {"SPEARMAN"}),
]


def _add_analysis_section(document, analysis_results: list[dict]) -> None:
    add_h2(document, "Resultados de análisis estadístico")
    headers = ["Código", "Tipo", "N válido", "Valor", "Estadístico", "p", "Efecto"]
    widths = [26, 30, 20, 22, 26, 22, 32]

    def rows_for(results: list[dict]) -> list[list[str]]:
        return [
            [
                result.get("result_code") or "—",
                result.get("result_type") or "—",
                str(result["n_valid"]) if result.get("n_valid") is not None else "—",
                _format_number(result.get("numeric_value")),
                _format_number(result.get("statistic_value")),
                _format_number(result.get("p_value"), decimals=4),
                _format_number(result.get("effect_size")),
            ]
            for result in results[:50]
        ]

    remaining = list(analysis_results)
    for label, result_types in _ANALYSIS_CATEGORIES:
        group = [r for r in remaining if r.get("result_type") in result_types]
        if not group:
            continue
        remaining = [r for r in remaining if r not in group]
        add_h3(document, label)
        add_colmena_table(document, headers, rows_for(group), widths_mm=widths, compact=True)

    if remaining:
        add_h3(document, "Otros resultados")
        add_colmena_table(document, headers, rows_for(remaining), widths_mm=widths, compact=True)


def _add_action_plan_section(document, number, plans: list[dict]) -> None:
    add_section_header(document, number, "Plan de acción y seguimiento")
    if not plans:
        add_body(document, "No hay planes de acción registrados para este estudio.")
        return
    for plan in plans:
        heading = add_h3(document, plan.get("name") or "Plan de acción")
        keep_with_next(heading)
        approval = "Aprobado" if plan.get("approved") else "Pendiente de aprobación"
        add_body(document, f"Estado: {plan.get('status') or '—'}. {approval}.")

        headers = ["Prioridad", "Constructo", "Hallazgo", "Medida", "Responsable", "Fecha", "Estado"]
        widths = [16, 28, 38, 36, 20, 18, 22]
        rows = []
        items = plan.get("items") or []
        for item in items:
            construct = item.get("construct_code") or "—"
            if item.get("construct_name"):
                construct += f" · {item['construct_name']}"
            rows.append([
                str(item.get("priority") or "—"),
                construct,
                item.get("finding") or "—",
                item.get("action_description") or "—",
                item.get("responsible_label") or "—",
                item.get("due_date") or "—",
                item.get("effective_status") or item.get("status") or "—",
            ])
        table = add_colmena_table(document, headers, rows, widths_mm=widths, compact=True)
        for row_index, item in enumerate(items):
            status_label = item.get("effective_status") or item.get("status") or "—"
            status = _ACTION_STATUS.get(status_label)
            if status:
                _render_status(table.rows[row_index + 1].cells[6], status_label, status)
        for item in items:
            for kpi in item.get("kpis") or []:
                measured = kpi.get("current_value")
                if measured is None:
                    latest = kpi.get("latest_measurement") or {}
                    measured = latest.get("text_value") or "sin medición"
                add_caption(
                    document,
                    f"KPI {kpi.get('code') or '—'}: {kpi.get('name') or '—'} · "
                    f"línea base {_format_number(kpi.get('baseline_value'))} · actual {measured} · "
                    f"meta {_format_number(kpi.get('target_value'))}.",
                )
        document.add_paragraph()


def _add_premium_section(document, premium: dict) -> None:
    add_h2(document, "Analítica premium")
    if premium.get("status") != "AVAILABLE":
        add_body(document, "No hay resultados premium publicables para este estudio.")
        return
    add_body(
        document,
        f"{premium.get('result_count', 0)} resultados publicables; "
        f"{premium.get('significant_result_count', 0)} con p ajustado o p menor a 0.05. "
        f"Métodos: {', '.join(premium.get('methods') or []) or '—'}.",
    )
    headers = ["Código", "Método", "N", "Estadístico", "p ajustado", "Efecto", "IC"]
    widths = [24, 26, 14, 24, 24, 34, 32]
    rows = []
    for result in (premium.get("results") or [])[:100]:
        adjusted = result.get("adjusted_p_value")
        rows.append([
            result.get("result_code") or "—",
            result.get("result_type") or "—",
            str(result.get("n_valid") if result.get("n_valid") is not None else "—"),
            _format_number(result.get("statistic_value")),
            _format_number(adjusted if adjusted is not None else result.get("p_value"), 4),
            " · ".join(filter(None, [_format_number(result.get("effect_size")), result.get("effect_label")])),
            f"[{_format_number(result.get('ci_lower'))}, {_format_number(result.get('ci_upper'))}]",
        ])
    add_colmena_table(document, headers, rows, widths_mm=widths, compact=True)
    add_caption(document, "Limitaciones: " + " ".join(premium.get("limitations") or []))


def _add_traceability_appendix(document, number, traceability: dict) -> None:
    add_section_header(document, number, "Anexo de trazabilidad")
    if not traceability:
        add_body(document, "No hay metadatos de trazabilidad disponibles.")
        return
    pairs = [
        ("Versión del instrumento", traceability.get("version_kind") or "—"),
        ("Hash del manifiesto", traceability.get("manifest_hash") or "—"),
        ("Linaje", " → ".join(traceability.get("lineage") or [])),
        ("N mínimo publicable", str((traceability.get("privacy") or {}).get("min_publishable_n", "—"))),
        ("Registros individuales incluidos", "No"),
    ]
    barem = traceability.get("barem") or {}
    if barem:
        pairs.extend([
            ("Baremo", f"{barem.get('name') or '—'} · {barem.get('version') or '—'}"),
            ("Hash del baremo", barem.get("content_hash") or "—"),
            ("Fuente del baremo", barem.get("source_reference") or "—"),
        ])
    add_label_value_table(document, pairs, widths_mm=[45, CONTENT_WIDTH_MM - 45])
    document.add_paragraph()

    runs = traceability.get("analysis_runs") or []
    if runs:
        add_h3(document, "Ejecuciones analíticas")
        headers = ["Tipo", "Estado", "Motor", "Algoritmo", "Hash de entrada"]
        widths = [30, 24, 40, 34, 50]
        rows = []
        for run in runs:
            rows.append([
                run.get("analysis_type") or "—",
                run.get("status") or "—",
                " ".join(filter(None, [run.get("engine"), run.get("engine_version")])) or "—",
                run.get("algorithm_version") or "—",
                run.get("input_hash") or "—",
            ])
        add_colmena_table(document, headers, rows, widths_mm=widths, compact=True)


def _add_signatures_section(document, number) -> None:
    # Firmas siempre en página propia — nunca queremos la grilla 2x2
    # partida entre dos páginas por quedar cerca del corte.
    document.add_page_break()
    add_section_header(document, number, "Firmas y conformidad")
    entries = [
        {"role": "Profesional responsable (Psicología/Salud ocupacional)"},
        {"role": "Representante de la organización"},
        {"role": "Seguridad y Salud en el Trabajo (SST)"},
        {"role": "Psicología ocupacional"},
    ]
    add_signature_grid(document, entries)


def _insert_chart(document, image: io.BytesIO | None) -> None:
    if image is None:
        return
    document.add_picture(image, width=Cm(15.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Formato
# ---------------------------------------------------------------------------

def _format_number(value: float | None, decimals: int = 2) -> str:
    if value is None:
        return "—"
    return f"{value:.{decimals}f}"


def _format_pct(value: float | None, decimals: int = 1) -> str:
    if value is None:
        return "—"
    return f"{value:.{decimals}f}%"


def _format_percentage(ratio: float | None, decimals: int = 1) -> str:
    if ratio is None:
        return "—"
    return f"{ratio * 100:.{decimals}f}%"


def _format_duration(seconds: float | None) -> str:
    if seconds is None:
        return "—"
    minutes, remainder = divmod(int(seconds), 60)
    return f"{minutes} min {remainder} s" if minutes else f"{remainder} s"


def _format_datetime(value: str | None) -> str | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value).strftime("%d/%m/%Y %H:%M")
    except ValueError:
        return value
