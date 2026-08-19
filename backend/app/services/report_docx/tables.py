"""Motor de tablas Colmena — nunca un `table.style = "Light Grid Accent 1"`.
Controla ancho, padding, tipografía, fondo, bordes finos y repetición de
encabezado en un único punto para que ninguna sección tenga que tocar OOXML
directamente."""

from __future__ import annotations

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .borders import set_cell_margins, set_table_borders
from .cells import set_cell_vertical_center, set_column_widths_mm, set_table_fixed_layout, shade_cell
from .theme import COLMENA, FONT

_HAIRLINE = {"sz": 4, "color": COLMENA["line"], "val": "single"}


def _set_repeat_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_mark = OxmlElement("w:tblHeader")
    header_mark.set(qn("w:val"), "true")
    tr_pr.append(header_mark)


def _style_cell_text(cell, *, size: float, bold: bool, color: str, align=None) -> None:
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def add_colmena_table(
    document,
    headers: list[str],
    rows: list[list[str]],
    widths_mm: list[float],
    *,
    header: bool = True,
    stripe: bool = True,
    compact: bool = False,
):
    total_width = sum(widths_mm)
    n_cols = len(widths_mm)
    table = document.add_table(rows=1 if header else 0, cols=n_cols)
    set_table_fixed_layout(table, total_width)

    pad = 55 if compact else 80

    if header:
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = text
            shade_cell(cell, COLMENA["navy"])
            set_cell_margins(cell, top=pad, bottom=pad, start=100, end=100)
            set_cell_vertical_center(cell)
            _style_cell_text(cell, size=FONT["table_header"], bold=True, color=COLMENA["white"])
        _set_repeat_header(table)

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        bg = COLMENA["surface_alt"] if (stripe and row_index % 2 == 1) else COLMENA["white"]
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = str(value)
            shade_cell(cell, bg)
            set_cell_margins(cell, top=pad, bottom=pad, start=100, end=100)
            set_cell_vertical_center(cell)
            _style_cell_text(cell, size=FONT["table"], bold=False, color=COLMENA["text"])

    set_table_borders(table, top=_HAIRLINE, bottom=_HAIRLINE, left=_HAIRLINE, right=_HAIRLINE)
    set_column_widths_mm(table, widths_mm)
    return table


def add_key_value_grid(document, pairs: list[tuple[str, str]], widths_mm: list[float]):
    """Grid de 4 columnas label/valor/label/valor para bloques tipo ficha
    técnica — las celdas de label van semibold, las de valor con más ancho."""
    n_rows = (len(pairs) + 1) // 2
    total_width = sum(widths_mm)
    table = document.add_table(rows=n_rows, cols=4)
    set_table_fixed_layout(table, total_width)

    for row_index in range(n_rows):
        row = table.rows[row_index]
        for col in range(2):
            pair_index = row_index * 2 + col
            label_cell = row.cells[col * 2]
            value_cell = row.cells[col * 2 + 1]
            bg = COLMENA["surface_alt"] if row_index % 2 == 1 else COLMENA["white"]
            if pair_index < len(pairs):
                label_text, value_text = pairs[pair_index]
            else:
                label_text, value_text = "", ""
            label_cell.text = label_text
            value_cell.text = value_text
            for cell, bold in ((label_cell, True), (value_cell, False)):
                shade_cell(cell, bg)
                set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
                set_cell_vertical_center(cell)
                _style_cell_text(cell, size=FONT["table"], bold=bold, color=COLMENA["text"])

    set_table_borders(table, top=_HAIRLINE, bottom=_HAIRLINE, left=_HAIRLINE, right=_HAIRLINE)
    set_column_widths_mm(table, widths_mm)
    return table


def add_label_value_table(document, pairs: list[tuple[str, str]], widths_mm: list[float]):
    """Tabla de 2 columnas etiqueta/valor a todo el ancho — para bloques con
    valores largos (hashes, linaje) donde la grilla de 4 columnas comprime
    demasiado."""
    total_width = sum(widths_mm)
    table = document.add_table(rows=0, cols=2)
    set_table_fixed_layout(table, total_width)

    for row_index, (label_text, value_text) in enumerate(pairs):
        cells = table.add_row().cells
        bg = COLMENA["surface_alt"] if row_index % 2 == 1 else COLMENA["white"]
        cells[0].text = label_text
        cells[1].text = value_text
        for cell, bold in ((cells[0], True), (cells[1], False)):
            shade_cell(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
            set_cell_vertical_center(cell)
            _style_cell_text(cell, size=FONT["table"], bold=bold, color=COLMENA["text"])

    set_table_borders(table, top=_HAIRLINE, bottom=_HAIRLINE, left=_HAIRLINE, right=_HAIRLINE)
    set_column_widths_mm(table, widths_mm)
    return table
