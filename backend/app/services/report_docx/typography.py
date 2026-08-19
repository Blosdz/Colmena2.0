"""Estilos de párrafo propios del documento — nunca se usa
`document.add_heading()` ni los estilos internos de Word (Heading 1/2, Light
Grid/List Accent 1) como apariencia final."""

from __future__ import annotations

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .theme import COLMENA, FONT, FONT_FAMILY

_STYLE_SPECS = {
    "Colmena Title": dict(size=FONT["cover_title"], bold=True, color=COLMENA["navy"], space_after=2),
    "Colmena Org": dict(size=FONT["cover_org"], bold=True, color=COLMENA["teal"], space_after=4),
    "Colmena Subtitle": dict(size=FONT["cover_subtitle"], color=COLMENA["muted"], space_after=10),
    "Colmena H1": dict(size=FONT["h1"], bold=True, color=COLMENA["navy"], space_after=2),
    "Colmena H2": dict(size=FONT["h2"], bold=True, color=COLMENA["navy"], space_before=10, space_after=6),
    "Colmena H3": dict(size=FONT["h3"], bold=True, color=COLMENA["navy"], space_before=8, space_after=4),
    "Colmena Body": dict(size=FONT["body"], color=COLMENA["text"], space_after=6),
    "Colmena Caption": dict(size=FONT["caption"], color=COLMENA["muted"], space_after=4),
    "Colmena Muted": dict(size=FONT["body"], color=COLMENA["muted"], space_after=4),
    "Colmena Label": dict(size=FONT["label"], bold=True, color=COLMENA["navy"], space_after=2),
}


def register_styles(document) -> None:
    styles = document.styles
    for name, spec in _STYLE_SPECS.items():
        style = styles[name] if name in [s.name for s in styles] else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = FONT_FAMILY
        font.size = Pt(spec["size"])
        font.bold = spec.get("bold", False)
        font.color.rgb = RGBColor.from_string(spec["color"])
        style.paragraph_format.space_before = Pt(spec.get("space_before", 0))
        style.paragraph_format.space_after = Pt(spec.get("space_after", 4))


def add_cover_title(document, text: str):
    p = document.add_paragraph(text, style="Colmena Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_cover_org(document, text: str):
    return document.add_paragraph(text, style="Colmena Org")


def add_cover_subtitle(document, text: str):
    return document.add_paragraph(text, style="Colmena Subtitle")


def add_h1(document, text: str):
    return document.add_paragraph(text, style="Colmena H1")


def add_h2(document, text: str):
    return document.add_paragraph(text, style="Colmena H2")


def add_h3(document, text: str):
    return document.add_paragraph(text, style="Colmena H3")


def add_body(document, text: str):
    return document.add_paragraph(text, style="Colmena Body")


def add_caption(document, text: str):
    return document.add_paragraph(text, style="Colmena Caption")


def add_muted(document, text: str):
    return document.add_paragraph(text, style="Colmena Muted")


def add_label(document, text: str):
    return document.add_paragraph(text, style="Colmena Label")
