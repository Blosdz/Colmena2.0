"""Fase 12 (cierre MVP del reporte CENSOPAS): portada, ficha técnica, perfil
sociolaboral, D1-D6/S1-S20 separados, interpretación, priorización, plan
preventivo, conclusiones, trazabilidad y paridad DOCX/PDF/JSON."""

import json
from io import BytesIO

from docx import Document
from httpx import AsyncClient

from app.tests.integration.test_censopas_flow import _setup_censopas_study

# El PDF (matplotlib) embebe el texto vía glyphs de una fuente subseteada,
# no como literales ASCII en el content stream — ni siquiera con
# `pdf.compression = 0` el texto queda buscable como substring sin un
# parser de PDF completo (pypdf/pdfplumber, no instalados). Por eso estos
# tests verifican el PDF solo estructuralmente (header válido, tamaño
# variando con el contenido) y verifican paridad de contenido real
# comparando DOCX vs JSON — ambos consumen exactamente las mismas
# funciones del bundle (`_trichotomy_lines` etc.) que alimentan el PDF,
# así que un DOCX/JSON correctos son evidencia fuerte de que el PDF
# (mismo bundle, mismo builder de líneas) también lo es. Ver reporte de
# auditoría POINT-5-REPORTS para la limitación documentada.


async def _add_sociolaboral_question(client: AsyncClient, version_id: int, code: str):
    return (
        await client.post(
            f"/api/v1/instrument-versions/{version_id}/items",
            json={
                "code": code,
                "question_text": "Eres",
                "question_type": "SINGLE_CHOICE",
                "research_role": "EXOGENOUS",
            },
        )
    ).json()


async def _generate(client: AsyncClient, study_id: int, output_format: str, sections=None):
    payload = {"output_format": output_format}
    if sections is not None:
        payload["sections"] = sections
    resp = await client.post(f"/api/v1/studies/{study_id}/reports", json=payload)
    assert resp.status_code == 201, resp.text
    download = await client.get(f"/api/v1/reports/{resp.json()['id']}/download")
    assert download.status_code == 200
    return download


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs) + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )


async def test_short_shows_dimensions_without_subdimensions(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    assert bundle["dimension_results"], "SHORT debe tener resultados D1-D6"
    assert bundle["subdimension_results"] == []

    docx_download = await _generate(client, study["id"], "DOCX")
    text = _docx_text(docx_download.content)
    assert "Resultados por dimensión (D1-D6)" in text
    assert "Resultados por subdimensión (S1-S20)" not in text


async def test_medium_shows_dimension_and_subdimension_in_separate_blocks(
    client: AsyncClient, seed_user, seed_project
) -> None:
    instrument = (
        await client.post("/api/v1/instruments", json={"name": "CENSOPAS media", "is_system": False})
    ).json()
    version = (
        await client.post(
            f"/api/v1/instruments/{instrument['id']}/versions",
            json={
                "version_code": "V1",
                "status": "DRAFT",
                "config": {
                    "censopas_version_kind": "MEDIUM",
                    "censopas_expected": {
                        "questions": 2,
                        "scored": 2,
                        "descriptive": 0,
                        "dimensions": 1,
                        "subdimensions": 1,
                    },
                },
            },
        )
    ).json()
    root = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/structure-variables",
            json={"code": "CENSOPAS", "name": "CENSOPAS", "role": "OUTCOME"},
        )
    ).json()

    d1_item = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/items",
            json={"code": "P1", "question_text": "D1 item", "question_type": "LIKERT"},
        )
    ).json()
    s1_item = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/items",
            json={"code": "P2", "question_text": "S1 item", "question_type": "LIKERT"},
        )
    ).json()

    dimension = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/constructs",
            json={"parent_id": root["id"], "code": "D1", "name": "Dimensión 1", "construct_type": "DIMENSION"},
        )
    ).json()
    subdimension = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/constructs",
            json={
                "parent_id": dimension["id"],
                "code": "S1",
                "name": "Subdimensión 1",
                "construct_type": "SUBDIMENSION",
            },
        )
    ).json()
    await client.post(
        f"/api/v1/constructs/{dimension['id']}/items",
        json={"question_id": d1_item["id"], "weight": 1, "scoring_direction": "DIRECT"},
    )
    await client.post(
        f"/api/v1/constructs/{subdimension['id']}/items",
        json={"question_id": s1_item["id"], "weight": 1, "scoring_direction": "DIRECT"},
    )
    scale_map = {"1": 0, "2": 25, "3": 50, "4": 75, "5": 100}
    await client.post(
        f"/api/v1/instrument-versions/{version['id']}/scoring-rules",
        json={"question_id": d1_item["id"], "parameters": {"map": scale_map}, "status": "ACTIVE"},
    )
    await client.post(
        f"/api/v1/instrument-versions/{version['id']}/scoring-rules",
        json={"question_id": s1_item["id"], "parameters": {"map": scale_map}, "status": "ACTIVE"},
    )
    barem = (
        await client.post(
            f"/api/v1/instrument-versions/{version['id']}/barems", json={"name": "Barem media"}
        )
    ).json()
    for construct_id in (dimension["id"], subdimension["id"]):
        await client.post(
            f"/api/v1/barems/{barem['id']}/cutoffs",
            json={"construct_id": construct_id, "cut_1": 33, "cut_2": 66, "direction": "LOWER_BETTER"},
        )

    survey = (
        await client.post(
            f"/api/v1/projects/{seed_project.id}/surveys/from-instrument",
            json={
                "created_by_user_id": seed_user.id,
                "instrument_version_id": version["id"],
                "name": "Encuesta media",
            },
        )
    ).json()
    study = (
        await client.post(
            f"/api/v1/projects/{seed_project.id}/studies",
            json={
                "survey_id": survey["id"],
                "name": "Estudio media",
                "study_type": "CENSO",
                "min_publishable_n": 1,
            },
        )
    ).json()
    await client.patch(f"/api/v1/studies/{study['id']}", json={"barem_id": barem["id"]})
    await client.post(f"/api/v1/studies/{study['id']}/open")
    for p1, p2 in [("5", "5"), ("4", "4"), ("2", "2")]:
        rs = (await client.post(f"/api/v1/studies/{study['id']}/response-sessions")).json()
        await client.put(f"/api/v1/response-sessions/{rs['id']}/responses/{d1_item['id']}", json={"raw_code": p1})
        await client.put(f"/api/v1/response-sessions/{rs['id']}/responses/{s1_item['id']}", json={"raw_code": p2})
        await client.post(f"/api/v1/response-sessions/{rs['id']}/complete")

    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    assert bundle["cover"]["version_kind"] == "MEDIUM"
    assert bundle["dimension_results"], "MEDIUM debe tener D1-D6"
    assert bundle["subdimension_results"], "MEDIUM debe tener S1-S20"

    docx_download = await _generate(client, study["id"], "DOCX")
    text = _docx_text(docx_download.content)
    assert "Resultados por dimensión (D1-D6)" in text
    assert "Resultados por subdimensión (S1-S20)" in text

    pdf_download = await _generate(client, study["id"], "PDF")
    assert pdf_download.content[:5] == b"%PDF-"


async def test_sociolaboral_profile_n_and_percentage_and_privacy(
    client: AsyncClient, seed_user, seed_project
) -> None:
    # Setup minimal propio: agregar el ítem EXOGENOUS antes de abrir el
    # estudio (una vez OPEN, la versión queda bloqueada para edición
    # estructural — InstrumentEditPolicy). No necesita scoring/dimensión.
    instrument = (
        await client.post("/api/v1/instruments", json={"name": "CENSOPAS sociolaboral", "is_system": False})
    ).json()
    version = (
        await client.post(
            f"/api/v1/instruments/{instrument['id']}/versions",
            json={"version_code": "V1", "status": "DRAFT"},
        )
    ).json()
    sex_item = await _add_sociolaboral_question(client, version["id"], "C-001")
    survey = (
        await client.post(
            f"/api/v1/projects/{seed_project.id}/surveys/from-instrument",
            json={
                "created_by_user_id": seed_user.id,
                "instrument_version_id": version["id"],
                "name": "Encuesta sociolaboral",
            },
        )
    ).json()
    study = (
        await client.post(
            f"/api/v1/projects/{seed_project.id}/studies",
            json={
                "survey_id": survey["id"],
                "name": "Estudio sociolaboral",
                "study_type": "CENSO",
                "min_publishable_n": 2,
            },
        )
    ).json()
    await client.post(f"/api/v1/studies/{study['id']}/open")

    # 3 "M" (publicable) y 1 "F" (suprimido, n=1 < min_publishable_n=2).
    for code in ("M", "M", "M", "F"):
        rs = (await client.post(f"/api/v1/studies/{study['id']}/response-sessions")).json()
        await client.put(
            f"/api/v1/response-sessions/{rs['id']}/responses/{sex_item['id']}", json={"raw_code": code}
        )
        await client.post(f"/api/v1/response-sessions/{rs['id']}/complete")

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    profile = bundle["sociolaboral_profile"]
    assert profile, "debe existir el perfil sociolaboral"
    question = next(q for q in profile if q["code"] == "C-001")
    by_label = {c["label"]: c for c in question["categories"]}
    assert by_label["M"]["n"] == 3
    assert by_label["M"]["suppressed"] is False
    assert by_label["F"]["suppressed"] is True
    assert by_label["F"]["n"] is None
    assert by_label["F"]["percentage"] is None

    docx_download = await _generate(client, study["id"], "DOCX")
    docx_text = _docx_text(docx_download.content)
    assert "Perfil sociolaboral" in docx_text
    assert "Oculto" in docx_text

    pdf_download = await _generate(client, study["id"], "PDF")
    assert pdf_download.content[:5] == b"%PDF-"


async def test_reference_barem_never_labeled_official(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    assert bundle["methodological_label"]["official_equivalence"] is False
    assert bundle["cover"]["methodological_label"]["official_equivalence"] is False

    docx_download = await _generate(client, study["id"], "DOCX")
    text = _docx_text(docx_download.content)
    assert "No equivalente al resultado oficial CENSOPAS-COPSOQ" in text
    assert "Reporte oficial CENSOPAS-COPSOQ" not in text


async def test_cover_includes_instrument_and_version(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    assert bundle["cover"]["instrument_label"] == "CENSOPAS-COPSOQ"
    assert bundle["cover"]["version_kind"] in {"SHORT", "MEDIUM", "UNKNOWN"}
    assert bundle["cover"]["study_code"] == f"EST-{study['id']:06d}"
    assert bundle["cover"]["confidentiality_notice"]


async def test_conclusions_present(client: AsyncClient, seed_user, seed_project) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    assert bundle["conclusions"]

    docx_download = await _generate(client, study["id"], "DOCX")
    text = _docx_text(docx_download.content)
    assert "Conclusiones" in text


async def test_plan_preventivo_fields_in_docx_pdf_json(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    plan = (
        await client.post(
            f"/api/v1/studies/{study['id']}/action-plans", json={"name": "Plan preventivo"}
        )
    ).json()
    item = (
        await client.post(
            f"/api/v1/action-plans/{plan['id']}/items",
            json={
                "title": "Reducir exposición",
                "finding": "Prioridad preventiva D1",
                "origin_hypothesis": "Carga elevada, plazos cortos.",
                "action_description": "Redistribuir la carga semanal.",
                "responsible_label": "Comité SST",
                "due_date": "2020-01-01",
                "construct_id": dimension["id"],
            },
        )
    ).json()
    await client.post(
        f"/api/v1/studies/{study['id']}/kpis",
        json={"action_plan_item_id": item["id"], "name": "% asistencia", "target_value": 90, "unit": "%"},
    )
    kpis = (await client.get(f"/api/v1/studies/{study['id']}/kpis")).json()
    await client.post(
        f"/api/v1/kpis/{kpis[0]['id']}/measurements",
        json={"measured_at": "2026-06-01T00:00:00Z", "numeric_value": 75},
    )

    json_download = await _generate(client, study["id"], "JSON")
    bundle = json.loads(json_download.text)
    action = bundle["action_plans"][0]["items"][0]
    assert action["origin_hypothesis"] == "Carga elevada, plazos cortos."
    assert action["effective_status"] == "OVERDUE"
    assert action["kpis"][0]["current_value"] == 75.0

    docx_text = _docx_text((await _generate(client, study["id"], "DOCX")).content)
    assert "Carga elevada" in docx_text
    assert "75" in docx_text
    pdf_download = await _generate(client, study["id"], "PDF")
    assert pdf_download.content[:5] == b"%PDF-"


async def test_docx_pdf_json_share_methodological_sections(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=1)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    docx_text = _docx_text((await _generate(client, study["id"], "DOCX")).content)
    pdf_download = await _generate(client, study["id"], "PDF")
    json_bundle = json.loads((await _generate(client, study["id"], "JSON")).text)

    assert pdf_download.content[:5] == b"%PDF-"
    for phrase in ["Perfil sociolaboral", "Priorización preventiva", "Conclusiones"]:
        assert phrase in docx_text, f"falta '{phrase}' en DOCX"
    for key in ["sociolaboral_profile", "priority_ranking", "conclusions", "interpretation", "cover"]:
        assert key in json_bundle
    for key in ["priority_ranking", "conclusions", "interpretation", "cover"]:
        assert json_bundle[key], f"'{key}' está vacío en JSON"


async def test_suppressed_dimension_never_leaks_in_any_format(
    client: AsyncClient, seed_user, seed_project
) -> None:
    study, _dimension = await _setup_censopas_study(client, seed_user, seed_project, min_publishable_n=5)
    scoring = await client.post(f"/api/v1/studies/{study['id']}/censopas/scoring")
    assert scoring.status_code == 200, scoring.text

    json_bundle = json.loads((await _generate(client, study["id"], "JSON")).text)
    result = json_bundle["dimension_results"][0]
    assert result["suppressed"] is True
    assert result["favorable_pct"] is None
    assert result["unfavorable_pct"] is None

    docx_text = _docx_text((await _generate(client, study["id"], "DOCX")).content)
    assert "Oculto" in docx_text

    pdf_download = await _generate(client, study["id"], "PDF")
    assert pdf_download.content[:5] == b"%PDF-"
