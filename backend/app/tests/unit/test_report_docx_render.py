"""Smoke test del design system DOCX: el bundle mínimo debe renderizar un
.docx válido y cada tabla debe llevar un único `w:tblW`/`w:tblLayout` — la
regresión concreta que rompió el ancho de las tablas del header (LibreOffice
fusiona dos `w:tbl` consecutivas sin párrafo entre medio, y un `tblPr` con
`w:tblW` duplicado hace que el renderer ignore el ancho fijado)."""

import io

from docx import Document
from docx.oxml.ns import qn

from app.services.report_docx import render_report_docx

_BUNDLE = {
    "study": {"name": "Estudio de prueba", "status": "OPEN", "study_type": "CENSO"},
    "cover": {"version_kind": "SHORT", "methodological_label": {"label": "Baremo de referencia"}},
    "sections": [],
    "traceability": {"privacy": {"min_publishable_n": 5}},
    "dimension_results": [
        {
            "construct_code": "D1",
            "construct_name": "Exigencias psicológicas",
            "favorable_pct": 20.0,
            "favorable_n": 4,
            "intermediate_pct": 30.0,
            "intermediate_n": 6,
            "unfavorable_pct": 50.0,
            "unfavorable_n": 10,
            "n_valid": 20,
            "collective_classification": "RIESGO_ALTO",
        }
    ],
}


def test_render_report_docx_produces_valid_document():
    docx_bytes = render_report_docx(_BUNDLE)
    document = Document(io.BytesIO(docx_bytes))
    assert len(document.paragraphs) > 0


def test_every_table_has_a_single_tblw_and_tbllayout():
    docx_bytes = render_report_docx(_BUNDLE)
    document = Document(io.BytesIO(docx_bytes))

    def check_tables(tables):
        for table in tables:
            tblPr = table._tbl.tblPr
            assert len(tblPr.findall(qn("w:tblW"))) <= 1, "tblW duplicado rompe el ancho en LibreOffice"
            assert len(tblPr.findall(qn("w:tblLayout"))) <= 1

    check_tables(document.tables)
    section = document.sections[0]
    check_tables(section.header.tables)
    check_tables(section.footer.tables)


if __name__ == "__main__":
    test_render_report_docx_produces_valid_document()
    test_every_table_has_a_single_tblw_and_tbllayout()
    print("ok")
